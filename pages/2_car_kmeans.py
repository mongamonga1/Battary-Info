# -*- coding: utf-8 -*-
"""차명별 K-means 군집 분석 (k 자동선정, 가로 스크롤, 저가 GPT 요약/Word Export - Chat Completions)"""
import warnings
warnings.filterwarnings("ignore")

import os, base64, hashlib
from io import BytesIO
from pathlib import Path
from math import pi

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib as mpl
import seaborn as sns

from sklearn.preprocessing import StandardScaler, OneHotEncoder
from sklearn.compose import ColumnTransformer
from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score
from sklearn.decomposition import PCA
from sklearn.manifold import TSNE

# ── 경량 테마(색상) ──────────────────────────────────────────────
def apply_colors(page_bg="#F5F7FB", sidebar_bg="#0F172A", sidebar_text="#DBE4FF", sidebar_link="#93C5FD"):
    st.markdown(f"""
    <style>
      .stApp {{ background: {page_bg}; }}
      section[data-testid="stSidebar"] {{ background: {sidebar_bg}; }}
      section[data-testid="stSidebar"] * {{ color: {sidebar_text} !important; }}
      section[data-testid="stSidebar"] a, section[data-testid="stSidebar"] svg {{
        color: {sidebar_link} !important; fill: {sidebar_link} !important;
      }}
      section[data-testid="stSidebar"] a:hover {{ background-color: rgba(255,255,255,0.08) !important; border-radius: 8px; }}
    </style>
    """, unsafe_allow_html=True)

apply_colors(
    page_bg="#F5F7FB",
    sidebar_bg="#0F172A",
    sidebar_text="#FFFFFF",
    sidebar_link="#93C5FD"
)
st.markdown("""
<style>
section[data-testid="stSidebar"] [data-testid*="FileUploaderDropzone"]{
  background-color:#1E293B !important; border:1.5px dashed #94A3B8 !important; border-radius:12px !important;
}
section[data-testid="stSidebar"] .stFileUploader [data-testid*="FileUploaderDropzone"],
section[data-testid="stSidebar"] .stFileUploader > div > div{
  background-color:#1E293B !important; border:1.5px dashed #94A3B8 !important; border-radius:12px !important;
}
section[data-testid="stSidebar"] [data-testid*="FileUploaderDropzone"] *:not(button):not([role="button"]):not(button *):not([role="button"] *),
section[data-testid="stSidebar"] .stFileUploader [data-testid*="FileUploaderDropzone"] *:not(button):not([role="button"]):not(button *):not([role="button"] *){
  color:#EAF2FF !important; opacity:1 !important; filter:none !important;
}
section[data-testid="stSidebar"] [data-testid*="FileUploader"] button,
section[data-testid="stSidebar"] [data-testid*="FileUploader"] [role="button"],
section[data-testid="stSidebar"] [data-testid*="FileUploader"] button *,
section[data-testid="stSidebar"] [data-testid*="FileUploader"] [role="button"] *{
  background-color:#F1F5F9 !important; color:#0F172A !important; font-weight:700 !important; opacity:1 !important;
}
section[data-testid="stSidebar"] [data-testid="stSelectbox"] [data-baseweb="select"] *{ color:#0F172A !important; }
div[data-baseweb="popover"] [data-baseweb="menu"] *{ color:#0F172A !important; }
</style>
""", unsafe_allow_html=True)

# ───────────────────────────── OpenAI secrets 헬퍼 ─────────────────────────────
def get_openai_conf():
    api_key = None
    model_name = None
    if hasattr(st, "secrets") and "openai" in st.secrets:
        sect = st.secrets["openai"]
        api_key = sect.get("api_key") or api_key
        model_name = sect.get("model") or model_name
    if hasattr(st, "secrets"):
        api_key = api_key or st.secrets.get("OPENAI_API_KEY")
    api_key = api_key or os.environ.get("OPENAI_API_KEY")
    return api_key, model_name

# ───────────────────────────── 선택 라이브러리 ─────────────────────────────
try:
    from scipy.cluster.hierarchy import linkage
    _has_scipy = True
except Exception:
    _has_scipy = False

try:
    from yellowbrick.cluster import KElbowVisualizer
    _has_yb = True
except Exception:
    _has_yb = False

# ───────────────────────────── Word Export 유틸 ─────────────────────────────
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

def _apply_korean_fonts(doc, font_name="Malgun Gothic", size_pt=11):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    rpr = style._element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    for k in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rFonts.set(qn(k), font_name)
    for h in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        if h in doc.styles:
            s = doc.styles[h]
            s.font.name = font_name
            rpr = s._element.get_or_add_rPr()
            rFonts = rpr.get_or_add_rFonts()
            for k in ("w:eastAsia", "w:ascii", "w:hAnsi"):
                rFonts.set(qn(k), font_name)

def _add_step(doc, n, title):
    p = doc.add_paragraph(f"Step {n}: {title}")
    p.style = "Heading 1"

def export_word_like_full(doc_title, model, gpt_analysis_text, main_imgs, profile_imgs,
                          dfm, num_pool, votes, k_final, font_name="Malgun Gothic"):
    doc = Document()
    _apply_korean_fonts(doc, font_name=font_name, size_pt=11)
    doc.add_heading(f"EV Battery Clustering Report – {model}", level=0)

    _add_step(doc, 1, "Data Loading & Preprocessing")
    doc.add_paragraph("✅ 1단계 완료: 데이터 로드 및 전처리 완료.")
    _add_step(doc, 2, "Model Selection & Filtering")
    doc.add_paragraph(f"✅ 2단계 완료: 선택된 모델 – {model}, 샘플 수 – {len(dfm):,}개")

    _add_step(doc, 3, "Optimal k Determination")
    k_sil  = votes.get("silhouette", "—")
    k_elb  = votes.get("elbow", "—")
    k_dend = votes.get("dendrogram", "—")
    doc.add_paragraph(f"Silhouette 최적 k: {k_sil}")
    doc.add_paragraph(f"Elbow 최적 k: {k_elb}")
    doc.add_paragraph(f"Dendrogram 추정 k: {k_dend}")
    doc.add_paragraph(f"✅ 3단계 완료: 최종 k = {k_final}")

    _add_step(doc, 4, "K-Means Clustering")
    counts = dfm["cluster"].value_counts().sort_index()
    doc.add_paragraph("✅ 4단계 완료: " + " / ".join([f"Cluster {c} → {int(counts[c])}개" for c in counts.index]))

    _add_step(doc, 5, "Visualizations")
    viz_idx = 1
    def add_img(cap, png_bytes):
        nonlocal viz_idx
        doc.add_paragraph(f"5-{viz_idx}: {cap}")
        doc.add_picture(BytesIO(png_bytes), width=Inches(6.2))
        viz_idx += 1
    for cap, png in main_imgs:    add_img(cap, png)
    for cap, png in profile_imgs: add_img(cap, png)

    _add_step(doc, 6, "Cluster-wise Summary")
    for para in str(gpt_analysis_text).split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_paragraph("✅ 6단계 완료: 클러스터별 통계 요약 완료.")

    means = dfm.groupby("cluster")[num_pool].mean().round(2)
    tbl = doc.add_table(rows=1, cols=2 + len(num_pool))
    hdr = tbl.rows[0].cells
    hdr[0].text = "Cluster"; hdr[1].text = "Count"
    for i, c in enumerate(num_pool, start=2): hdr[i].text = f"Mean {c}"
    for c in counts.index:
        row = tbl.add_row().cells
        row[0].text = str(c); row[1].text = str(int(counts[c]))
        for j, col in enumerate(num_pool, start=2):
            row[j].text = str(means.loc[c, col])

    bio = BytesIO(); doc.save(bio); bio.seek(0)
    return bio

# ───────────────────────────── 기본 설정 ─────────────────────────────
mpl.rcParams["font.family"] = "DejaVu Sans"
mpl.rcParams["axes.unicode_minus"] = False
st.header("🚗 차명별 K-means 군집 분석")

# ───────────────────────────── 데이터 로드 ─────────────────────────────
DATA_PATH = Path("data/SoH_NCM_Dataset_selected_Fid_및_배터리등급열추가.xlsx")

@st.cache_data(show_spinner=False)
def load_excel(path_or_buffer) -> pd.DataFrame:
    df = pd.read_excel(path_or_buffer, engine="openpyxl")
    df.columns = df.columns.map(lambda x: str(x).strip())
    return df

uploaded_file = st.sidebar.file_uploader("엑셀 업로드(선택)", type=["xlsx"])

df_raw = None
if uploaded_file is not None:
    df_raw = load_excel(uploaded_file)
    st.success("업로드한 파일을 사용합니다.")
elif DATA_PATH.exists():
    df_raw = load_excel(DATA_PATH)

if df_raw is None:
    st.error("기본 엑셀 파일을 찾을 수 없습니다. 사이드바에서 업로드해 주세요.")
    st.stop()

# ───────────────────────────── 컬럼 표준화 ─────────────────────────────
@st.cache_data(show_spinner=False)
def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    def pick_first(cands):
        for c in cands:
            if c in out.columns: return c
        return None
    mapping = {}
    schema = [
        ("Model",       ["차명", "배터리종류", "차종", "모델"]),
        ("Age",         ["사용연수(t)", "사용연수", "연식"]),
        ("SoH",         ["SoH_pred(%)", "SoH(%)", "SOH"]),
        ("Price",       ["중고거래가격", "개당가격", "거래금액", "가격"]),
        ("CellBalance", ["셀 간 균형", "셀간균형"]),
    ]
    for std, cands in schema:
        src = pick_first(cands)
        if src: mapping[src] = std
    out = out.rename(columns=mapping)
    if out.columns.duplicated().any():
        out = out.loc[:, ~out.columns.duplicated()]
    if "CellBalance" in out.columns:
        out["CellBalance"] = (
            out["CellBalance"]
            .map({"우수":"Good","정상":"Normal","경고":"Warning","심각":"Critical"})
            .fillna(out["CellBalance"])
        )
    if "Price" in out.columns:
        out["Price"] = (out["Price"].astype(str).str.replace(r"[^\d.\-]", "", regex=True)
                        .pipe(pd.to_numeric, errors="coerce"))
    if "Age" in out.columns:
        out["Age"] = pd.to_numeric(out["Age"], errors="coerce")
    if "SoH" in out.columns:
        out["SoH"] = pd.to_numeric(out["SoH"], errors="coerce")
    return out

df = normalize_columns(df_raw)

# 필수/수치 컬럼 확인
if "Model" not in df.columns:
    st.error("엑셀에 '차명/배터리종류/차종/모델' 중 하나가 없어 Model 컬럼이 없습니다.")
    st.stop()

num_pool = [c for c in ["Age","SoH","Price"] if c in df.columns]
if len(num_pool) < 2:
    st.error(f"수치 컬럼이 부족합니다(필요≥2). 현재: {num_pool}")
    st.stop()

# ───────────────────────────── 사이드바 컨트롤 ─────────────────────────────
models        = sorted(df["Model"].dropna().astype(str).unique())
choice        = st.sidebar.selectbox("차명 선택", models)
show_tsne     = st.sidebar.checkbox("t-SNE 2D 추가", value=True)
show_pca3     = st.sidebar.checkbox("PCA 3D 추가", value=False)
perplexity    = st.sidebar.slider("t-SNE perplexity", 5, 50, 30, 1)
show_profiles = st.sidebar.checkbox("추가 프로파일(가로 스크롤)", value=True)

# 💸 비용 옵션 (최소 과금 구조)
st.sidebar.markdown("### 💸 비용 옵션")
cost_saver   = st.sidebar.checkbox("비용 절감 모드(저가 모델·짧은 응답)", value=True)
DEFAULT_MODEL = "gpt-4o-mini"
_api_key, _model_from_secret = get_openai_conf()
MODEL_NAME   = _model_from_secret or DEFAULT_MODEL
# GPT 요약 내부 설정
cost_saver  = False
TEMPERATURE = 0.2
MAX_TOKENS  = 320 if cost_saver else 600

# ───────────────────────────── 모델 데이터 준비 ─────────────────────────────
sub_all = df[df["Model"].astype(str) == str(choice)].copy().dropna(subset=num_pool)
n = len(sub_all)
if n < 3:
    st.warning(f"'{choice}' 유효 표본이 {n}건이라 분석할 수 없습니다(≥3 필요).")
    st.stop()

ks = list(range(2, min(10, n)))
preproc = ColumnTransformer(
    transformers=[
        ("num", StandardScaler(), num_pool),
        ("cat", OneHotEncoder(drop="first", handle_unknown="ignore"),
         ["CellBalance"] if "CellBalance" in sub_all.columns else []),
    ],
    remainder="drop",
)
X = preproc.fit_transform(sub_all)
if hasattr(X, "toarray"): X = X.toarray()

# ───────────────────────────── k 자동선정 (속도개선版) ─────────────────────────────
@st.cache_data(show_spinner=False)
def choose_k_multi_fast(X: np.ndarray, ks: tuple, sample_size: int = 1000, rnd: int = 42, use_scipy: bool = False):
    """한 루프에서 inertia/labels를 얻어 Silhouette과 Elbow를 동시 계산. Silhouette은 표본샘플."""
    sil_map, inertia_list = {}, []
    rng = np.random.RandomState(rnd)

    for k in ks:
        if k >= len(X):
            break
        km = KMeans(n_clusters=k, random_state=rnd, n_init=10)
        labels = km.fit_predict(X)
        inertia_list.append((k, km.inertia_))
        # 실루엣: 전체가 크면 표본샘플로
        try:
            s = silhouette_score(
                X, labels,
                sample_size=min(sample_size, len(X) - 1) if len(X) > 50 else None,
                random_state=rnd
            )
            sil_map[k] = s
        except Exception:
            pass

    # Silhouette 최댓값
    k_sil = max(sil_map, key=sil_map.get) if sil_map else None

    # Elbow: inertia 차분의 최대 변화점
    k_vals, inert = zip(*inertia_list) if inertia_list else ([], [])
    k_elb = None
    if len(inert) >= 2:
        diffs = np.diff(inert)
        idx = int(np.argmax(diffs))
        k_elb = k_vals[min(idx + 1, len(k_vals)-1)]

    # Dendrogram(옵션, 표본 최대 200)
    k_dend = None
    if use_scipy:
        try:
            m = X.shape[0]
            idx = np.arange(m if m <= 200 else 200)
            Z = linkage(X[idx], method="ward")
            dists = Z[:, 2]
            gaps = np.diff(dists)
            if len(gaps) >= 1:
                k_est = m - (int(np.argmax(gaps)) + 1)
                k_dend = max(2, min(k_est, ks[-1]))
        except Exception:
            pass

    votes = {"silhouette": k_sil, "elbow": k_elb}
    if k_dend is not None:
        votes["dendrogram"] = k_dend

    valid = [v for v in (k_sil, k_elb, k_dend) if v is not None]
    k_final = int(np.median(valid)) if valid else 3
    return k_final, votes

k_final, votes = choose_k_multi_fast(X, tuple(ks), sample_size=1000, rnd=42, use_scipy=_has_scipy)
st.caption(f"선택된 k = {k_final} (Sil={votes.get('silhouette','—')}, "
           f"Elbow={votes.get('elbow','—')}, Dend={votes.get('dendrogram','—')} → median)")

# ───────────────────────────── 최종 학습 & 라벨 (캐싱) ─────────────────────────────
@st.cache_data(show_spinner=False)
def fit_kmeans_labels(X: np.ndarray, k: int, rnd: int = 42):
    km = KMeans(n_clusters=k, random_state=rnd, n_init=10)
    return km.fit_predict(X)

labels = fit_kmeans_labels(X, k_final, rnd=42)
sub_all = sub_all.copy(); sub_all["cluster"] = labels
clusters = sorted(sub_all["cluster"].unique())

# ───────────────────────────── 유틸: fig → png/base64 ─────────────────────────────
def fig_to_png(fig, dpi=160):
    buf = BytesIO(); fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight"); plt.close(fig)
    return buf.getvalue()
def to_b64(png_bytes): return base64.b64encode(png_bytes).decode("utf-8")

# ───────────────────────────── 공통 CSS(가로 스크롤) ─────────────────────────────
st.markdown("""
<style>
.scroll-x { overflow-x:auto; padding:8px 0 10px; }
.scroll-row { display:inline-flex; gap:16px; }
.scroll-row img { border-radius:12px; box-shadow:0 2px 8px rgba(0,0,0,.12); }
.caption-center { text-align:center; color:#6b7280; font-size:12px; }
</style>
""", unsafe_allow_html=True)

# ───────────────────────────── 결과 그래프(가로 스크롤) ─────────────────────────────
main_images = []

# PCA 2D
p2 = PCA(2, random_state=42).fit_transform(X)
fig = plt.figure(figsize=(5.2, 4.0))
plt.scatter(p2[:,0], p2[:,1], c=labels, cmap="tab10", s=55, edgecolors="k", alpha=0.9)
plt.title(f"{choice}: PCA 2D (k={k_final})"); plt.xlabel("PC1"); plt.ylabel("PC2"); plt.tight_layout()
main_images.append(("PCA 2D", fig_to_png(fig)))

# Radar
mean_matrix = sub_all.groupby("cluster")[num_pool].mean()
norm_means = mean_matrix.copy()
for c in num_pool:
    mn, mx = df[c].min(), df[c].max()
    norm_means[c] = 0.5 if (pd.isna(mn) or pd.isna(mx) or mx==mn) else (norm_means[c]-mn)/(mx-mn)
angles = [i/len(num_pool)*2*pi for i in range(len(num_pool))] + [0]
fig = plt.figure(figsize=(5.2, 4.0)); ax = plt.subplot(111, polar=True)
for i in clusters:
    vals = norm_means.loc[i].tolist() + [norm_means.loc[i].tolist()[0]]
    ax.plot(angles, vals, label=f"Cluster {i}"); ax.fill(angles, vals, alpha=0.1)
ax.set_xticks(angles[:-1]); ax.set_xticklabels(num_pool)
plt.title(f"{choice}: Radar (k={k_final})"); plt.legend(loc="upper right", bbox_to_anchor=(1.25,1.05))
main_images.append(("Radar", fig_to_png(fig)))

# t-SNE 2D (캐싱 + 자동 다운샘플)
@st.cache_data(show_spinner=False)
def compute_tsne(X: np.ndarray, perplexity: int, rnd: int = 42, max_n: int = 2000):
    # 큰 데이터는 다운샘플링해서 임베딩(시각화 용도) → 속도↑
    if X.shape[0] > max_n:
        idx = np.random.RandomState(rnd).choice(X.shape[0], max_n, replace=False)
        X_use = X[idx]
        idx_out = idx
    else:
        X_use = X
        idx_out = None
    ts2 = TSNE(n_components=2, perplexity=min(perplexity, X_use.shape[0]-1),
               init="pca", max_iter=500, random_state=rnd).fit_transform(X_use)
    return ts2, idx_out

if show_tsne:
    ts2, idx_out = compute_tsne(X, perplexity, rnd=42, max_n=2000)
    lab_plot = labels if idx_out is None else labels[idx_out]
    fig = plt.figure(figsize=(5.2, 4.0))
    plt.scatter(ts2[:,0], ts2[:,1], c=lab_plot, cmap="tab10", s=55, edgecolors="k", alpha=0.9)
    plt.title(f"{choice}: t-SNE 2D (k={k_final})"); plt.xlabel("t-SNE1"); plt.ylabel("t-SNE2"); plt.tight_layout()
    main_images.append(("t-SNE 2D", fig_to_png(fig)))

# PCA 3D (옵션)
if show_pca3:
    from mpl_toolkits.mplot3d import Axes3D  # noqa: F401
    p3 = PCA(3, random_state=42).fit_transform(X)
    fig = plt.figure(figsize=(5.6, 4.2)); ax3 = fig.add_subplot(111, projection="3d")
    ax3.scatter(p3[:,0], p3[:,1], p3[:,2], c=labels, cmap="tab10", s=50, edgecolors="k", alpha=0.85)
    ax3.set_title(f"{choice}: PCA 3D (k={k_final})"); ax3.set_xlabel("PC1"); ax3.set_ylabel("PC2"); ax3.set_zlabel("PC3")
    main_images.append(("PCA 3D", fig_to_png(fig)))

# 화면 출력(가로 스크롤)
html_imgs = "".join([f"<img src='data:image/png;base64,{base64.b64encode(p).decode('utf-8')}' height='320'/>" for _,p in main_images])
st.markdown(f"<div class='scroll-x'><div class='scroll-row'>{html_imgs}</div></div>", unsafe_allow_html=True)
st.markdown("<div class='caption-center'>좌우 스크롤로 결과 그래프(PCA2D, Radar, 옵션: t-SNE/PCA3D)를 확인하세요.</div>", unsafe_allow_html=True)

# ───────────────────────────── 추가 프로파일(가로 스크롤) ─────────────────────────────
profile_images = []
if show_profiles:
    for col in num_pool:
        fig = plt.figure(figsize=(6,4)); sns.boxplot(x="cluster", y=col, data=sub_all, palette="tab10")
        plt.title(f"{choice}: {col} by Cluster (k={k_final})")
        profile_images.append((f"Box {col}", fig_to_png(fig)))
    if "CellBalance" in sub_all.columns:
        fig = plt.figure(figsize=(6,4))
        sns.countplot(x="cluster", hue="CellBalance", data=sub_all, palette="Set2")
        plt.title(f"{choice}: Count of CellBalance by Cluster")
        profile_images.append(("Count CellBalance", fig_to_png(fig)))

        ctab_pct = pd.crosstab(sub_all["cluster"], sub_all["CellBalance"], normalize="index")*100
        ctab_pct = ctab_pct.reindex(clusters, fill_value=0)
        fig = plt.figure(figsize=(6,4)); ax = plt.gca()
        ctab_pct.plot(kind="bar", stacked=True, colormap="Paired", ax=ax)
        plt.title(f"{choice}: CellBalance Distribution (%) by Cluster"); plt.tight_layout()
        profile_images.append(("Stacked CellBalance", fig_to_png(fig)))

    mean_matrix = sub_all.groupby("cluster")[num_pool].mean()
    fig = plt.figure(figsize=(6,4))
    sns.heatmap(mean_matrix, annot=True, cmap="coolwarm", fmt=".2f")
    plt.title(f"{choice}: Numeric Feature Means per Cluster")
    profile_images.append(("Heatmap Means", fig_to_png(fig)))

    html_prof = "".join([f"<img src='data:image/png;base64,{base64.b64encode(p).decode('utf-8')}' height='300'/>" for _,p in profile_images])
    st.markdown(f"<div class='scroll-x'><div class='scroll-row'>{html_prof}</div></div>", unsafe_allow_html=True)
    st.markdown("<div class='caption-center'>추가 프로파일도 가로 스크롤로 확인하세요.</div>", unsafe_allow_html=True)

# ───────────────────────────── GPT 요약 & Word 내보내기 ─────────────────────────────
st.subheader("🧠 믿:음 분석결과 & Word 분석보고서")

if "ai_text" not in st.session_state:
    st.session_state.ai_text = None

def summarize_compact(dfm: pd.DataFrame, num_pool: list[str]) -> str:
    counts = dfm["cluster"].value_counts().sort_index()
    means  = dfm.groupby("cluster")[num_pool].mean()
    line_counts = f"N={len(dfm)}, k={dfm['cluster'].nunique()}"
    line_means  = " | ".join([
        "C{}: ".format(c) + ", ".join([f"{col} {means.loc[c,col]:.1f}" for col in num_pool])
        for c in counts.index
    ])
    return line_counts + "\n" + line_means

def generate_ai_summary(model, k_final, votes, dfm, num_pool, model_name, max_tokens, temperature):
    try:
        from openai import OpenAI
        _has_openai = True
    except Exception:
        _has_openai = False

    stats_compact = summarize_compact(dfm, num_pool)
    try:
        if not _has_openai:
            raise RuntimeError("openai 패키지가 설치되어 있지 않습니다.")
        api_key, model_from_secret = get_openai_conf()
        if model_from_secret: model_name = model_from_secret
        if not api_key: raise RuntimeError("OPENAI_API_KEY not set")

        client = OpenAI(api_key=api_key)
        system_msg = (
            "You are a concise Korean data analyst. "
            "군집분석 결과를 250~350자 한국어 본문으로 요약하라. "
            "군집별 (연식·SoH·가격) 비교와 실무 활용 포인트 2~3개 포함. "
            "불필요한 표/이모지/목록은 지양."
        )
        user_prompt = (
            f"[모델]{model}\n"
            f"[최종 k] {k_final} (Sil={votes.get('silhouette')}, "
            f"Elbow={votes.get('elbow')}, Dend={votes.get('dendrogram')})\n"
            f"[요약통계]\n{stats_compact}"
        )
        resp = client.chat.completions.create(
            model=model_name,
            messages=[{"role":"system","content":system_msg},
                      {"role":"user","content":user_prompt}],
            temperature=temperature,
            max_tokens=max_tokens,
        )
        return resp.choices[0].message.content.strip()
    except Exception:
        cluster_means = dfm.groupby("cluster")[num_pool].mean().round(1)
        top_price = cluster_means["Price"].idxmax() if "Price" in cluster_means.columns else "—"
        return (f"[로컬 요약] {model}을(를) k={k_final}로 군집화했습니다. "
                f"SoH·연식·가격 평균 기준 군집 간 차이가 확인됩니다. "
                f"SoH·가격이 높은 군집({top_price})은 리마케팅 타깃, "
                f"저SoH 군집은 정밀 점검 권고가 유효합니다.")

col_a, col_b = st.columns([1,2])
with col_a:
    gen_btn = st.button("🧠 분석결과 생성 & Word로 저장", use_container_width=True)
with col_b:
    if st.session_state.ai_text:
        st.markdown("**🔎 분석 결과 (믿:음 생성)**")
        st.write(st.session_state.ai_text)

if gen_btn:
    with st.spinner("믿:음 분석결과 생성 및 Word 문서 작성 중..."):
        ai_text = generate_ai_summary(
            model=choice, k_final=k_final, votes=votes, dfm=sub_all,
            num_pool=num_pool, model_name=MODEL_NAME,
            max_tokens=MAX_TOKENS, temperature=TEMPERATURE
        )
        st.session_state.ai_text = ai_text

        word_buf = export_word_like_full(
            doc_title=f"EV 배터리 군집 분석 보고서 – {choice}",
            model=choice, gpt_analysis_text=ai_text,
            main_imgs=main_images, profile_imgs=profile_images if show_profiles else [],
            dfm=sub_all, num_pool=num_pool, votes=votes, k_final=k_final,
            font_name="Malgun Gothic"
        )

    st.success("보고서를 생성했습니다.")
    st.download_button(
        "⬇️ Word 파일 다운로드",
        data=word_buf,
        file_name=f"EV_Battery_Report_{choice}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
